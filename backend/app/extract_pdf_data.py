import fitz
from docx import Document
import io

## the layout logic (multilayout and single layout will be implemented in future with llm/visual model)
def extract_pdf_data(file_bytes: bytes) -> dict:
    doc = fitz.open(stream=file_bytes, filetype="pdf")

    text_parts = []
    has_images = False
    is_multi_column = False

    for page in doc:
        text_parts.append(page.get_text())

        if page.get_images():
            has_images = True

        words = page.get_text("words")
        if len(words) > 20:
            page_width = page.rect.width
            mid = page_width / 2
            lines: dict[int, list] = {}
            for w in words:
                line_no = w[6]
                lines.setdefault(line_no, []).append(w)

            # multi_column_lines = 0
            # for line_words in lines.values():
            #     left_words = [w for w in line_words if w[0] < mid - 20]
            #     right_words = [w for w in line_words if w[0] > mid + 20]
            #     if len(left_words) >= 3 and len(right_words) >= 3:
            #         multi_column_lines += 1

            # if multi_column_lines >= 4:
            #     is_multi_column = True

    has_tables = False
    for page in doc:
        try:
            if page.find_tables().tables:
                has_tables = True
                break
        except Exception:
            pass

    return {
        "text": "\n".join(text_parts).strip(),
        "has_tables": has_tables,
        "has_images": has_images,
        "is_multi_column": False,
    }

def extract_docx_data(file_bytes: bytes) -> dict:
    doc = Document(io.BytesIO(file_bytes))

    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
    has_tables = len(doc.tables) > 0

    has_images = False
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            has_images = True
            break

    return {
        "text": "\n".join(text_parts).strip(),
        "has_tables": has_tables,
        "has_images": has_images,
        "is_multi_column": False,  
    }


def extract_document_data(filename: str, file_bytes: bytes) -> dict:
    if filename.lower().endswith(".pdf"):
        return extract_pdf_data(file_bytes)
    elif filename.lower().endswith(".docx"):
        return extract_docx_data(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}")
